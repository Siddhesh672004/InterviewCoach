"""Resume and job description text extraction.

Supports PDF (pypdf), DOCX (python-docx), and plain text.
Both extractors are tolerant of malformed files — they return the best-effort
text and never raise; on hard failure they return an empty string and the
caller decides how to surface the error.
"""
from __future__ import annotations
import io
import re
from typing import BinaryIO


MAX_RESUME_CHARS = 12_000  # ~3-4 LLM pages, plenty for any normal resume
MAX_JD_CHARS = 8_000


def _normalize(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(stream: BinaryIO) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(stream)
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        return _normalize("\n\n".join(pages))
    except Exception:
        return ""


def extract_docx(stream: BinaryIO) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return ""
    try:
        document = docx.Document(stream)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return _normalize("\n".join(parts))
    except Exception:
        return ""


def extract_resume(filename: str, content: bytes) -> str:
    """Dispatch on file extension. Returns trimmed text or empty string."""
    if not content:
        return ""
    name = (filename or "").lower()
    stream = io.BytesIO(content)

    if name.endswith(".pdf"):
        text = extract_pdf(stream)
    elif name.endswith(".docx"):
        text = extract_docx(stream)
    elif name.endswith(".doc"):
        # Old binary .doc format is not supported by python-docx; fall back
        # to a naive text decode so users at least get something.
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
        text = _normalize(text)
    elif name.endswith(".txt"):
        text = _normalize(content.decode("utf-8", errors="ignore"))
    else:
        # Try PDF first, then DOCX, then plain text — last-ditch dispatch
        text = extract_pdf(io.BytesIO(content)) or extract_docx(io.BytesIO(content))
        if not text:
            text = _normalize(content.decode("utf-8", errors="ignore"))

    return text[:MAX_RESUME_CHARS]


def clean_jd(text: str) -> str:
    return _normalize(text or "")[:MAX_JD_CHARS]
